from loguru import logger
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from zipfile import ZipFile
import os

from . import shops_parser
from . import parse_ids
from .classes.shops import *
from .exceptions import *


logger.add("debug.log")

# os.getcwd() needed to get current directory depending on process that imported this module
TMP_FOLDER = f"{os.getcwd()}/tmp/"
OUTPUT_FOLDER = f"{os.getcwd()}/output/"

EMPTY_FILENAME = "1. Пустой.docx"
EMPTY_ID = "0"
EMPTY_SHOP = Shop(EMPTY_ID, "", "")

OUTPUT_ZIP_FILENAME = "0. Накладные.zip"
OUTPUT_NAMES = [
    "2. Вторник",
    "3. Среда",
    "4. Четверг",
    "5. Пятница",
    "6. Суббота",
    "7. Воскресенье",
]


class InvoiceGenerator:

    def __init__(self,
                 input_ids_filename: str,
                 input_shops_filename: str,
                 template_filename: str,
                 ids_separator: str = "TestName",
                 output_names: list[str] = OUTPUT_NAMES,
                 add_empty_file: bool = True,
                 start_tmp_count_from: int = 1):

        # self.ids = ids_string.split(ids_separator)
        # self.ids_string = ids_string
        # self.ids_separator = ids_separator
        
        self.ids_by_days: list[list[str]] = parse_ids.parse_ids(input_ids_filename)
        
        self.template_filename = template_filename
        self.output_names = output_names
        self.output_name: str = ""
        self.start_tmp_count_from = start_tmp_count_from
        self.shops = InvoiceGenerator.parse_shops(input_shops_filename)
        self._input_shops_filename = input_shops_filename

        # Existing files analysis can be here
        self.invoices_filenames: list[str] = []
        self.add_empty_file = add_empty_file

    @staticmethod
    def parse_shops(input_shops_filename: str) -> list[Shop]:
        try:
            return shops_parser.parse_shops(input_shops_filename)
        except FileNotFoundError as err:
            message = f"No such file in directory: {input_shops_filename}"
            logger.error(message)
            raise err
        except OSError as err:
            message = f"""Check input file name type (must be string). Current filename: 
                - name: {input_shops_filename}
                - type: {type(input_shops_filename)}"""
            logger.error(message)
            raise err
        except Exception as err:
            logger.error(err)
            raise err

    def generate_all(self, zip: bool = True) -> None:
        for day_index, name in enumerate(OUTPUT_NAMES):
            self.invoices_filenames = []
            self.output_name = name
            ids: list[str] = self.ids_by_days[day_index]
            self.generate_file(ids)
            
        if self.add_empty_file:
            self.generate_empty_file()
        
        self.zip_output()
        

    def zip_output(self):
        zip_object = ZipFile(OUTPUT_FOLDER + OUTPUT_ZIP_FILENAME, "w")
        for filename in OUTPUT_NAMES:
            zip_object.write(f"{OUTPUT_FOLDER}{filename}.docx")
        zip_object.close()

    def generate_empty_file(self) -> None:
        self.generate_separate_invoice_file(EMPTY_FILENAME, EMPTY_ID, path=OUTPUT_FOLDER)
        ...
    
    def generate_file(self, ids: list[str], do_all: bool = True) -> None:
        """
        Generates final merged file. Can do only merging or also call other functions needed to do all process.

            - do_all: [bool] Need to repair.
        """
        
        if do_all is True:
            self.generate_separate_invoice_files(ids)
        else:
            err_message = "do_all argument functionality is not created yet. Put it on True (default value)"
            logger.error(err_message)
            raise NotCreatedFunctionalityError(err_message)

        err_message = "There is no invoice files generated. You can turn do_all to True to do it automatically"
        try:
            assert len(self.invoices_filenames) > 0, err_message
        except AssertionError as err:
            logger.error(err_message)
            raise err

        if len(self.invoices_filenames) == 1:
            self._rename_alone_docx()
            logger.info(f"There was only one invoice file {self.invoices_filenames[0]}. It was renamed {self.output_name}.docx")
            return
        else:
            self._merge_invoice_files()

    def generate_separate_invoice_files(self, ids) -> None:
        count = self.start_tmp_count_from
        for shop_id in ids:
            # Creating filename
            invoice_filename = f"{self.output_name}_{count}.docx"

            # Appending it to filenames stack
            self.invoices_filenames.append(invoice_filename)

            self.generate_separate_invoice_file(invoice_filename, shop_id)
            count += 1

    def generate_separate_invoice_file(self, invoice_filename: str, shop_id: str, path: str = TMP_FOLDER) -> None:
        doc = DocxTemplate(self.template_filename)
        if shop_id == EMPTY_ID:
            shop = EMPTY_SHOP
        else:
            shop = self._find_shop(shop_id)
        context = self._generate_context(shop)
        doc.render(context)
        doc.save(f"{path}{invoice_filename}")

        logger.info(f"Invoice file generated (Shop info: {shop})")

    def _merge_invoice_files(self):
        target_name = self.output_name

        master = Document(f"{TMP_FOLDER}{self.invoices_filenames[0]}")
        composer = Composer(master)
        for filename in self.invoices_filenames[1:]:
            doc = Document(f"{TMP_FOLDER}{filename}")
            composer.append(doc)
        composer.save(f"{OUTPUT_FOLDER}{target_name}.docx")
        logger.info(f"Merged file generated: {target_name}.docx")

    def _rename_alone_docx(self):
        filename = self.invoices_filenames[0]
        doc = Document(filename)
        doc.save(self.output_name + ".docx")

    def _find_shop(self, shop_id) -> Shop:
        for shop in self.shops:
            if shop.id == shop_id:
                return shop

        # If shop wasn't found
        message = f"Invalid ID: there are no shop with current ID ({shop_id}). Please check ID list or original TSV file"
        logger.error(message)
        raise InvalidShopIDError(message)

    @staticmethod
    def _generate_context(shop) -> dict[str, str]:
        context = {
            "organisation": shop.organisation,
            "address": shop.address
        }
        return context


if __name__ == "__main__":
    invoicer = InvoiceGenerator("target.md", "Магазины За ХЛЕБ! - Main info.tsv", "Template.docx")
    invoicer.generate_all()
    pass
